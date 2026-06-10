"""
ChromaDB 知识库工具 — 文件上传、语义检索、分块管理
"""
import os
import uuid
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_core.documents import Document

from src.hr_assistant.config import EMBEDDING_CONFIG

PERSIST_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))), "chroma_db")
COLLECTION_NAME = "faq_knowledge"

_embeddings = None
_vector_store = None

CHUNK_SIZE = 500
CHUNK_OVERLAP = 80


def _get_embeddings():
    global _embeddings
    if _embeddings is None:
        _embeddings = OpenAIEmbeddings(
            model=EMBEDDING_CONFIG["model"],
            api_key=EMBEDDING_CONFIG["api_key"],
            base_url=EMBEDDING_CONFIG["base_url"],
        )
    return _embeddings


def get_store():
    global _vector_store
    if _vector_store is None:
        _vector_store = Chroma(
            embedding_function=_get_embeddings(),
            persist_directory=PERSIST_DIR,
            collection_name=COLLECTION_NAME,
        )
    return _vector_store


def get_retriever(k: int = 3):
    return get_store().as_retriever(search_kwargs={"k": k})


def search(query: str, k: int = 3) -> list[Document]:
    return get_retriever(k).invoke(query)


def hybrid_search(query: str, k: int = 3) -> list[Document]:
    """混合检索：关键词提取 → 源筛选 → 内容评分排序"""
    import re
    
    # 关键词 → 文件名关键词映射
    KW_MAP = {
        "出差": ["出差休假"], "休假": ["出差休假"], "年假": ["出差休假"], "病假": ["出差休假"],
        "考勤": ["出差休假"], "请假": ["出差休假"], "加班": ["出差休假"], "住宿": ["出差休假"],
        "婚假": ["出差休假"], "产假": ["出差休假"], "丧假": ["出差休假"],
        "组织": ["组织机构"], "机构": ["组织机构"], "部门": ["组织机构"],
        "住房": ["住房保障"], "保障": ["住房保障"], "公寓": ["住房保障"],
        "招聘": ["招聘调配"], "调配": ["招聘调配"], "录用": ["招聘调配"], "面试": ["招聘调配"],
        "退出": ["招聘调配"], "入职": ["招聘调配"], "离职": ["招聘调配"],
        "年金": ["企业年金"], "企业年金": ["企业年金"],
        "教育": ["教育培训"], "培训": ["教育培训"], "学历": ["教育培训"],
        "绩效": ["绩效考核"], "考核": ["绩效考核"], "评分": ["绩效考核"],
        "公积金": ["公积金"], "公积": ["公积金"], "缴存": ["公积金"], "住房公积": ["公积金"],
        "出境": ["因私出境"], "出国": ["因私出境"], "护照": ["因私出境"], "因私": ["因私出境"],
        "干部": ["干部督查"], "监督": ["干部督查"], "选拔": ["干部督查"], "督查": ["干部督查"],
    }
    
    # 提取查询中的关键词：2-4 字中文 n-gram
    chinese_chars = re.findall(r'[\u4e00-\u9fff]', query)
    keywords = set()
    for n in (2, 3, 4):
        for i in range(len(chinese_chars) - n + 1):
            kw = ''.join(chinese_chars[i:i+n])
            keywords.add(kw)
    
    # 找到可能相关的源
    target_sources = set()
    for kw in keywords:
        if kw in KW_MAP:
            for src_key in KW_MAP[kw]:
                target_sources.add(src_key)
    
    # 从 store 获取候选文档
    store = get_store()
    all_results = store.get(include=["metadatas", "documents"])
    
    if not target_sources:
        # 无关键词匹配，回退到语义搜索
        return get_retriever(k).invoke(query)
    
    # 过滤：只保留 source 包含目标关键词的文档
    filtered = []
    for i, doc_id in enumerate(all_results.get("ids", [])):
        src = all_results["metadatas"][i].get("source", "")
        if any(tgt in src for tgt in target_sources):
            filtered.append(Document(
                page_content=all_results["documents"][i],
                metadata=all_results["metadatas"][i],
            ))
    
    if not filtered:
        return get_retriever(k).invoke(query)
    
    # 按关键词命中数 + 内容长度评分排序
    def score(doc):
        s = 0
        for kw in keywords:
            if kw in doc.page_content:
                s += 1
        s += min(len(doc.page_content) / 500.0, 1.0)
        return s
    
    ranked = sorted(filtered, key=score, reverse=True)
    return ranked[:k]


def rerank(query: str, docs: list[Document], top_n: int = 3) -> list[Document]:
    """BGE-Reranker-V2-M3 重排序：从候选文档中精选最相关的 top_n 个"""
    if len(docs) <= top_n:
        return docs

    try:
        from src.hr_assistant.config import RERANKER_CONFIG
        import requests

        url = RERANKER_CONFIG["base_url"].rstrip("/") + "/rerank"
        resp = requests.post(
            url,
            headers={
                "Authorization": f"Bearer {RERANKER_CONFIG['api_key']}",
                "Content-Type": "application/json",
            },
            json={
                "model": RERANKER_CONFIG["model"],
                "query": query,
                "documents": [d.page_content for d in docs],
            },
            timeout=30,
        )

        if resp.status_code != 200:
            return docs[:top_n]

        results = resp.json().get("results", [])
        if not results:
            return docs[:top_n]

        # 按 relevance_score 排序
        scored = sorted(results, key=lambda r: r.get("relevance_score", 0), reverse=True)
        ranked = [docs[r["index"]] for r in scored[:top_n] if r["index"] < len(docs)]
        return ranked or docs[:top_n]

    except Exception:
        return docs[:top_n]


def smart_search(query: str, k: int = 3) -> list[Document]:
    """智能检索：粗召回(10) → rerank → 精选(3)"""
    candidates = hybrid_search(query, k=10)
    return rerank(query, candidates, top_n=k)


def get_doc_count() -> int:
    try:
        return get_store()._collection.count()
    except Exception:
        return 0


# ──────────── 文本加载 ────────────

def _load_pdf(path: str) -> str:
    loader = PyPDFLoader(path)
    pages = loader.load()
    return "\n\n".join(p.page_content.strip() for p in pages if p.page_content.strip())


def _load_docx(path: str) -> str:
    try:
        from docx import Document as DocxDoc
        doc = DocxDoc(path)
        return "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    except Exception:
        raise ValueError("无法解析 .docx 文件，请确认格式正确")


def _load_file(path: str) -> tuple[str, str]:
    """返回 (text, filename)"""
    name = os.path.basename(path)
    ext = os.path.splitext(name)[1].lower()
    if ext == ".pdf":
        return _load_pdf(path), name
    elif ext == ".docx":
        return _load_docx(path), name
    else:
        # txt, md 等文本文件 — 直接用 open 读取
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read(), name
            except UnicodeDecodeError:
                continue
        raise ValueError(f"无法解码文件: {name}")


def _split_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", ". ", ",", " "],
    )
    return splitter.split_text(text)


# ──────────── 增/删/改 ────────────

def add_text(text: str, source: str) -> int:
    """添加文本到知识库（自动分块），返回块数"""
    chunks = _split_text(text)
    store = get_store()
    docs = [Document(page_content=c, metadata={"source": source, "chunk_id": str(uuid.uuid4())})
            for c in chunks]
    store.add_documents(docs)
    return len(chunks)


def add_file(path: str) -> tuple[str, int]:
    """上传文件到知识库，返回 (source_name, chunk_count)"""
    text, name = _load_file(path)
    count = add_text(text, source=name)
    return name, count


def add_file_content(name: str, content: str) -> int:
    """按文件名和内容添加"""
    return add_text(content, source=name)


def delete_source(source: str) -> int:
    """删除某个来源的所有文档块，返回删除数"""
    store = get_store()
    results = store.get(where={"source": source})
    ids = results.get("ids", [])
    if ids:
        store.delete(ids=ids)
    return len(ids)


def delete_chunk(chunk_id: str) -> bool:
    try:
        get_store().delete(ids=[chunk_id])
        return True
    except Exception:
        return False


def rechunk_source(source: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> int:
    """重新分块某个来源（先删后加），返回新块数"""
    store = get_store()
    results = store.get(where={"source": source}, include=["documents"])
    original_text = "\n\n".join(results.get("documents", []))
    if not original_text.strip():
        return 0
    delete_source(source)
    chunks = _split_text(original_text, chunk_size, overlap)
    docs = [Document(page_content=c, metadata={"source": source, "chunk_id": str(uuid.uuid4())}) for c in chunks]
    store.add_documents(docs)
    return len(chunks)


def rebuild_from_file(path: str) -> int:
    """从文件重建知识库（清空后重新导入）"""
    store = get_store()
    try:
        ids = store.get().get("ids", [])
        if ids:
            store.delete(ids=ids)
    except Exception:
        pass
    _, count = add_file(path)
    return count


# ──────────── 查/列 ────────────

def list_sources() -> list[dict]:
    """按来源分组统计"""
    store = get_store()
    results = store.get(include=["metadatas", "documents"])
    groups = {}
    for i, meta in enumerate(results.get("metadatas", [])):
        src = meta.get("source", "unknown")
        if src not in groups:
            groups[src] = {"source": src, "chunk_count": 0, "total_chars": 0, "sample": ""}
        groups[src]["chunk_count"] += 1
        groups[src]["total_chars"] += len(results["documents"][i])
        if not groups[src]["sample"]:
            groups[src]["sample"] = results["documents"][i][:100]
    return sorted(groups.values(), key=lambda x: x["total_chars"], reverse=True)


def list_chunks(source: str) -> list[dict]:
    """列出某来源的所有文档块"""
    store = get_store()
    results = store.get(where={"source": source}, include=["metadatas", "documents"])
    chunks = []
    for i, doc_id in enumerate(results.get("ids", [])):
        chunks.append({
            "id": doc_id,
            "source": results["metadatas"][i].get("source", ""),
            "chunk_id": results["metadatas"][i].get("chunk_id", ""),
            "content": results["documents"][i],
            "length": len(results["documents"][i]),
            "preview": results["documents"][i][:150] + ("..." if len(results["documents"][i]) > 150 else ""),
        })
    return chunks


def get_chunk_detail(chunk_id: str) -> dict | None:
    """获取单个块的完整内容"""
    store = get_store()
    results = store.get(ids=[chunk_id], include=["metadatas", "documents"])
    if not results.get("ids"):
        return None
    return {
        "id": chunk_id,
        "source": results["metadatas"][0].get("source", ""),
        "content": results["documents"][0],
        "length": len(results["documents"][0]),
    }
